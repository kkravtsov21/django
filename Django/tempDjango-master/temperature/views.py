from django.shortcuts import render
from django.http import HttpResponse
from django.shortcuts import render
from collections import defaultdict
from .models import Sensor, Temperature
from django.views.generic import CreateView
from .forms import SensorForm, TemperatureForm
from dicttoxml import dicttoxml
from xml.dom.minidom import parseString
from django.forms.models import model_to_dict
from docx import Document
from docx.shared import Cm
from django.shortcuts import redirect
from django.http import FileResponse


def index(request):
    return render(
        request, 'index.html',
    )

def locations(request):
    sensor = Sensor.objects.all()
    loc_list = []
    current_list = []

    for s in sensor:
        loc = s.location
        loc_list.append(loc)
        if loc not in current_list:
            current_list.append(loc)

    return render(
        request, 'locations.html', context={
            'loc_list':loc_list,
            'current_list':current_list
            },
    )

def temperatures(request):
    sensor = Sensor.objects.all()
    name_list = []
    loc_list = []
    id_list = []
    current_list = []

    for s in sensor:
        index = s.id
        id_list.append(index)

        name = s.name
        name_list.append(name)

        loc = s.location
        loc_list.append(loc)
        if loc not in current_list:
            current_list.append(loc)

    temperature = Temperature.objects.all()
    temp_list = []

    for t in temperature:
        temp = t.temperature
        temp_list.append(temp)

    sl = sensor.values('name').order_by('location')
    tl = temperature.values('id').order_by('sensor__location')

    return render(
        request, 'temperatures.html', context={
            'id_list':index,
            'name_list':name_list,
            'loc_list':loc_list,
            'temp_list':temp_list,
            'current_list':current_list,
            'sl':sl,
            't':t,
            'tl':tl,
        })

def sensorForm(request):
    if request.method == 'POST':
        form = SensorForm(request.POST)
        data = Sensor
        if form.is_valid():
            db = form.save()
            return render(request, 'success.html')
    else:
        form = SensorForm()
    return render(request, 'sensor_form.html', context={'form': form})

def temperatureForm(request):
    if request.method == 'POST':
        form = TemperatureForm(request.POST)
        if form.is_valid():
            db = form.save()
            return render(request, 'success.html')
    else:
        form = TemperatureForm()
    return render(request, 'temperature_form.html', context={'form': form})

def exports(request):
    return render(request, "exports.html")

def export_xml(request):
    try:
        exporter = Exporter()
        exporter.export_to_xml(get_data())
        return HttpResponse(open('data.xml').read(), content_type='text/xml')
    except Exception:
        return render(request, "fail.html", context={"data": 'xml'})

def export_docx(request):
    try:
        exporter = Exporter()
        exporter.export_to_docx(get_data())
        return FileResponse(open('data.docx', 'rb'))
    except Exception:
        return render(request, "fail.html", context={"data": 'docx'})

def success(request, data):
    return render(request, "success.html", context={"data": data})

class Exporter(object):

    def export_to_xml(self, data):
        xml = dicttoxml(data, custom_root='Data', attr_type=False)
        dom = parseString(xml)
        file = open('data.xml', 'w')
        file.write(dom.toprettyxml())
        file.close()

    def export_to_docx(self, data):
        document = Document()
        document.add_heading('Sensor Data', 0)
        document.add_paragraph(str(document.core_properties.created))
        for k, v in data.items():
            document.add_heading(k, level=1)
            for v1 in v:
                document.add_heading("Item", level=2).paragraph_format.left_indent = Cm(0.5)
                for v2 in v1:
                    document.add_paragraph(str(v2) + ": " + str(v1[v2]), style='List Bullet').paragraph_format.left_indent = Cm(1.5)

        document.save('data.docx')

class ModelData(object):
    def __init__(self):
        self.data = {}

    def add_data(self, data, name):
        self.__convert_to_dict(data, name)

    def set_data(self, data):
        self.data = data

    def get_data(self):
        return self.data

    def map(self, name_data1, name_data2, filed_cmp, field_fill):
        for d in self.data[name_data1]:
            d[filed_cmp] = next(d2 for d2 in self.data[name_data2] if d2["id"]==d[filed_cmp])[field_fill]

    def __convert_to_dict(self, data, name):
        cur_data = []
        for d in data:
            cur_data.append(model_to_dict(d))
        self.data[name] = cur_data

def get_data():
    data_model = ModelData()
    sensors = Sensor.objects.all()
    temperature = Temperature.objects.all()
    data_model.add_data(sensors, "Sensors")
    data_model.add_data(temperature, "Temperature")
    return data_model.get_data()